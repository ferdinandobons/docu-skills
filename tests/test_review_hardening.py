# SPDX-License-Identifier: MIT
"""Regression tests for the post-0.8.0 review hardening pass.

Each class pins one confirmed finding:

  * ``RoleIdValidationTest`` - ``schema.role_id`` validates its own composition,
    so the helper can never emit an id ``is_valid_role_id`` would reject (an
    invalid family/qualifier raises ``ValueError`` at the composition point
    instead of silently failing a later role lookup).
  * ``CaptionIndexBoundsTest`` - ``refresh_visible_caption_index_cache``
    validates BOTH field-inventory bounds before any ``children[...]`` access;
    a corrupted inventory (negative or past-the-end index) skips the index
    cleanly instead of corrupting the body via negative indexing or raising
    ``IndexError``.
  * ``ApplyBordersHardeningTest`` - ``_apply_borders`` refuses a parsed side
    element whose qualified tag is not the expected ``w:<side>`` and swallows
    only the two parse failures a profile string can cause.
  * ``SetOnlyWhenUnsetEmptyStringTest`` - the set-only-when-unset guards treat
    an explicitly EMPTY attribute (``w:val=""``, not valid OOXML) as unset, so
    a malformed template cannot pin a captured fact behind an empty string,
    while an authored value is still never clobbered.
"""

from __future__ import annotations

import sys
import unittest
from pathlib import Path
from unittest import mock

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "scripts"))

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn

from brandkit.formats.docx import generate as docx_generate
from brandkit.formats.docx import structure as docx_structure
from brandkit.profile import schema

_W_DECL = nsdecls("w")


class RoleIdValidationTest(unittest.TestCase):
    def test_valid_compositions_round_trip(self) -> None:
        self.assertEqual(schema.role_id("heading", 1), "heading.1")
        self.assertEqual(schema.role_id("list", "bullet", 2), "list.bullet.2")
        self.assertEqual(schema.role_id("paragraph"), "paragraph")
        self.assertEqual(
            schema.role_id("table", "default", "header"), "table.default.header"
        )

    def test_none_and_empty_qualifiers_are_dropped(self) -> None:
        self.assertEqual(schema.role_id("heading", None, ""), "heading")

    def test_composed_id_always_satisfies_is_valid_role_id(self) -> None:
        for rid in (
            schema.role_id("heading", 1),
            schema.role_id("list", "number", 3),
            schema.role_id("callout", "info"),
        ):
            self.assertTrue(schema.is_valid_role_id(rid))

    def test_empty_family_raises(self) -> None:
        with self.assertRaises(ValueError):
            schema.role_id("")

    def test_non_lowercase_family_raises(self) -> None:
        with self.assertRaises(ValueError):
            schema.role_id("Heading", 1)

    def test_special_characters_in_qualifier_raise(self) -> None:
        with self.assertRaises(ValueError):
            schema.role_id("paragraph", "Executive Summary")

    def test_dotted_injection_in_qualifier_raises_when_segment_invalid(self) -> None:
        # A qualifier carrying its own separator stays parseable only when every
        # resulting segment is valid; an empty segment must be refused.
        with self.assertRaises(ValueError):
            schema.role_id("list", "bullet.")


class CaptionIndexBoundsTest(unittest.TestCase):
    def _doc_with_paragraphs(self, texts: list[str]):
        doc = Document()
        for t in texts:
            doc.add_paragraph(t)
        return doc

    def _body_texts(self, doc) -> list[str]:
        return [p.text for p in doc.paragraphs]

    def _run_with_inventory(self, doc, inventory: list[dict]) -> int:
        with mock.patch.object(
            docx_structure, "_toc_field_begins", return_value=inventory
        ):
            return docx_structure.refresh_visible_caption_index_cache(
                doc, {"Tabella": ["Tabella 1. One", "Tabella 2. Two"]}
            )

    def test_negative_begin_index_skips_cleanly(self) -> None:
        # end >= begin holds (1 >= -2), so the inventory filter alone does not
        # protect this case: without the bounds check, children[-2] silently
        # grabbed the WRONG paragraph (python negative indexing) and the splice
        # corrupted the body.
        doc = self._doc_with_paragraphs(["a", "b", "c", "d"])
        before = self._body_texts(doc)
        inventory = [
            {
                "seq_id": "Tabella",
                "begin_index": -2,
                "end_index": 1,
                "instr": 'TOC \\c "Tabella"',
            }
        ]
        rebuilt = self._run_with_inventory(doc, inventory)
        self.assertEqual(rebuilt, 0)
        self.assertEqual(self._body_texts(doc), before)

    def test_past_the_end_indices_skip_cleanly(self) -> None:
        doc = self._doc_with_paragraphs(["a", "b"])
        before = self._body_texts(doc)
        inventory = [
            {
                "seq_id": "Tabella",
                "begin_index": 50,
                "end_index": 60,
                "instr": 'TOC \\c "Tabella"',
            }
        ]
        rebuilt = self._run_with_inventory(doc, inventory)
        self.assertEqual(rebuilt, 0)
        self.assertEqual(self._body_texts(doc), before)

    def test_no_entries_is_a_fast_noop(self) -> None:
        doc = self._doc_with_paragraphs(["a"])
        self.assertEqual(docx_structure.refresh_visible_caption_index_cache(doc, {}), 0)


class ApplyBordersHardeningTest(unittest.TestCase):
    def _fresh_ppr(self):
        return Document().add_paragraph("x")._p.get_or_add_pPr()

    def test_valid_side_is_appended_verbatim(self) -> None:
        ppr = self._fresh_ppr()
        serialized = f'<w:top {_W_DECL} w:val="single" w:sz="4"/>'
        docx_generate._apply_borders(ppr, {"top": serialized})
        pbdr = ppr.find(qn("w:pBdr"))
        self.assertIsNotNone(pbdr)
        top = pbdr.find(qn("w:top"))
        self.assertIsNotNone(top)
        self.assertEqual(top.get(qn("w:val")), "single")

    def test_mislabeled_side_is_refused(self) -> None:
        # A w:top element stored under the "bottom" key must never land under
        # w:pBdr as the wrong side.
        ppr = self._fresh_ppr()
        serialized = f'<w:top {_W_DECL} w:val="single"/>'
        docx_generate._apply_borders(ppr, {"bottom": serialized})
        pbdr = ppr.find(qn("w:pBdr"))
        self.assertIsNotNone(pbdr)
        self.assertIsNone(pbdr.find(qn("w:bottom")))
        self.assertIsNone(pbdr.find(qn("w:top")))

    def test_malformed_xml_is_skipped_without_crashing(self) -> None:
        ppr = self._fresh_ppr()
        docx_generate._apply_borders(ppr, {"top": "<<not-xml"})
        pbdr = ppr.find(qn("w:pBdr"))
        self.assertIsNotNone(pbdr)
        self.assertIsNone(pbdr.find(qn("w:top")))

    def test_encoding_declaration_string_is_skipped_without_crashing(self) -> None:
        # lxml raises ValueError (not XMLSyntaxError) for a unicode string that
        # carries an encoding declaration; the narrow catch must still skip it.
        ppr = self._fresh_ppr()
        serialized = f"<?xml version='1.0' encoding='UTF-8'?><w:top {_W_DECL}/>"
        docx_generate._apply_borders(ppr, {"top": serialized})
        pbdr = ppr.find(qn("w:pBdr"))
        self.assertIsNotNone(pbdr)
        self.assertIsNone(pbdr.find(qn("w:top")))

    def test_authored_side_is_never_clobbered(self) -> None:
        ppr = self._fresh_ppr()
        authored = f'<w:top {_W_DECL} w:val="double"/>'
        docx_generate._apply_borders(ppr, {"top": authored})
        captured = f'<w:top {_W_DECL} w:val="single"/>'
        docx_generate._apply_borders(ppr, {"top": captured})
        pbdr = ppr.find(qn("w:pBdr"))
        tops = pbdr.findall(qn("w:top"))
        self.assertEqual(len(tops), 1)
        self.assertEqual(tops[0].get(qn("w:val")), "double")


class SetOnlyWhenUnsetEmptyStringTest(unittest.TestCase):
    def test_absent_attribute_is_set(self) -> None:
        el = OxmlElement("w:spacing")
        docx_generate._set_twips_if_unset(el, "before", 240)
        self.assertEqual(el.get(qn("w:before")), "240")

    def test_empty_string_attribute_counts_as_unset(self) -> None:
        el = OxmlElement("w:spacing")
        el.set(qn("w:before"), "")
        docx_generate._set_twips_if_unset(el, "before", 240)
        self.assertEqual(el.get(qn("w:before")), "240")

    def test_authored_attribute_is_preserved(self) -> None:
        el = OxmlElement("w:spacing")
        el.set(qn("w:before"), "100")
        docx_generate._set_twips_if_unset(el, "before", 240)
        self.assertEqual(el.get(qn("w:before")), "100")

    def test_level_facts_empty_val_is_reasserted(self) -> None:
        lvl = OxmlElement("w:lvl")
        numfmt = OxmlElement("w:numFmt")
        numfmt.set(qn("w:val"), "")
        lvl.append(numfmt)
        docx_generate._reassert_level_facts(lvl, {"numFmt": "bullet"})
        self.assertEqual(numfmt.get(qn("w:val")), "bullet")

    def test_level_facts_authored_val_is_preserved(self) -> None:
        lvl = OxmlElement("w:lvl")
        numfmt = OxmlElement("w:numFmt")
        numfmt.set(qn("w:val"), "decimal")
        lvl.append(numfmt)
        docx_generate._reassert_level_facts(lvl, {"numFmt": "bullet"})
        self.assertEqual(numfmt.get(qn("w:val")), "decimal")

    def test_level_facts_empty_indent_attr_is_reasserted(self) -> None:
        lvl = OxmlElement("w:lvl")
        ppr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "")
        ind.set(qn("w:hanging"), "360")
        ppr.append(ind)
        lvl.append(ppr)
        docx_generate._reassert_level_facts(
            lvl, {"indent": {"left": 720, "hanging": 99}}
        )
        self.assertEqual(ind.get(qn("w:left")), "720")  # empty -> reasserted
        self.assertEqual(ind.get(qn("w:hanging")), "360")  # authored -> preserved


if __name__ == "__main__":
    unittest.main()
